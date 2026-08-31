"""Extracción de texto plano desde archivos de CV (pdf, docx, txt)."""
from pathlib import Path

from pypdf import PdfReader
from docx import Document


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Formato no soportado: {suffix}")


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    partes = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(partes)


def _extract_docx(path: Path) -> str:
    doc = Document(str(path))
    partes = [p.text for p in doc.paragraphs]
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    partes.append(celda.text)
    return "\n".join(partes)


FORMATOS_SOPORTADOS = {".pdf", ".docx", ".txt"}
